import docx
from docx import Document
from docx.shared import Pt, RGBColor
from news_nlp import get_ng_nlp_from_json, NewsItemNLP
import os
from news import json_dir, out_dir

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink, new_run

def apply_font(run_, rgb_color: RGBColor=None):
    run_.font.name = "Arial"
    run_.font.size = Pt(8)
    if not rgb_color:
        run_.font.color.rgb = RGBColor(0,0,0)
    else:
        run_.font.color.rgb = rgb_color

def convert_ng_nlp_iterable_to_docx(ni_iterable, docx_path):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(8)

    for ni in ni_iterable:
        # ni_nlp = NewsItemNLP() #remove when running
        p = document.add_paragraph()

        header = p.add_run(ni.header.strip() + ".")
        apply_font(header)
        header.font.color.rgb = RGBColor(255,134,134)
        header.bold = True
        
        p.add_run(" ")
        r = p.add_run(ni.summary)
        apply_font(r)
        p.add_run("\n")
        # r = p.add_run(ni_nlp.ni.url)
        hl, r = add_hyperlink(p, ni.url, ni.url, "7EB1FF", True)
        p.add_run("\n")

        p.style = document.styles['Normal']
    
    document.save(docx_path)
    print(f"saved {docx_path}")

def convert_ng_nlp_to_docx(ng_nlp_json_path, docx_path, append_filter=None):
    convert_ng_nlp_iterable_to_docx(
        [
            ni_nlp.ni for ni_nlp
            in get_ng_nlp_from_json(ng_nlp_json_path, append_filter=append_filter)
        ], docx_path
    )

docx_out = os.path.join(out_dir, "docx")
if not os.path.exists(docx_out):
    os.makedirs(docx_out)

if __name__ == "__main__":
    a = r"D:\dev local\news_summary\out\json\https___www_reuters_com_business_finance_.json"

    fa, fb = os.path.splitext(os.path.basename(a))

    b = os.path.join(json_dir, fa + "--nlp" + fb)
    b_applied = os.path.join(json_dir, fa + "--nlp--applied" + fb)

    docx_path = os.path.join(docx_out, fa + ".docx")
    
    convert_ng_nlp_to_docx(b_applied, docx_path)