from dataclasses import asdict, dataclass
import enum
from typing import List
from docx.api import Document
from news import NewsMinimal

@dataclass
class NewsDocx(NewsMinimal):
    category: str

def yield_news_from_docx(docx_path, date_str, enable_print=False):
    print("------ GETTING FROM DOCX {} ------".format(docx_path))
    doc = Document(docx_path)
    label = None
    found_1len_run_str = set()
    found_doc = set()
    found_header = set()
    for t in doc.tables:
        texts_list = [
            ''.join([cell.text for cell in row.cells])
            for row in t.rows
        ]
        texts = ''.join(texts_list)
        if "DISCLOSURE:" in texts:
            continue
        if "For details, please contact BAIPHIL " in texts:
            continue
        
        for row in t.rows:
            texts = [cell.text for cell in row.cells]
            text = ''.join(texts)
            if "Responsive and Sustainable Banking in the New Normal" in text:
                continue
            
            for cell in row.cells:
                if len(cell.text) == "":
                    continue
                for p in cell.paragraphs:
                    # print(p)
                    if len(p.runs) == 1:
                        if not cell.text in found_1len_run_str:
                            table_header = ' '.join(cell.text.split('\n'))
                            if table_header in ['PHILIPPINES', 'REST OF THE WORLD']:
                                if enable_print: print("!!!!!1len ]]]{}[[[".format(table_header))
                                if table_header == 'PHILIPPINES':
                                    label = "PHL"
                                else:
                                    label = "ROW"
                            found_1len_run_str.add(cell.text)
                    elif len(p.runs) > 1:
                        # print("-----------")
                        c = 1
                        header = p.runs[0:c]
                        runs = p.runs[c:]
                        while not '.' in ''.join([h.text for h in header]) and len(header) < len(p.runs):
                            c += 1
                            header = p.runs[0:c]
                            runs = p.runs[c:]
                        if len(header) == len(p.runs):
                            # print("!!! skipping {}".format(p.text))
                            pass
                        else:
                            header_text = ''.join([h.text for h in header]).strip()
                            if header_text[-1] != '.':
                                header_text = header_text + '.'
                            if header_text in found_header:
                                continue
                            if enable_print: print(header_text)
                            found_header.add(header_text)
                            valid_run_text = list()
                            for i, r in enumerate(runs):
                                if len(r.text) > 2:
                                    # print("{} ({}) >> {}".format(i, len(r.text),r.text))
                                    valid_run_text.append(r.text.strip()) 
                            summary = ' '.join(valid_run_text)
                            doc = header_text + " " + summary
                            if not doc in found_doc:
                                yield NewsDocx(None, header_text, summary, doc, date_str, label)
                                found_doc.add(doc)

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import linear_kernel
from news_nlp_lemma import tokenizer, token_stop
from news_nlp import preprocess
VECTORIZER: TfidfVectorizer = TfidfVectorizer(stop_words=token_stop, tokenizer=tokenizer)
def get_tfidf(documents):
    doc_vectors = VECTORIZER.fit_transform(documents)
    for idx, doc in enumerate(documents):
        cosine_similarities = linear_kernel(doc_vectors[idx:idx+1], doc_vectors).flatten()
        dscores_idx = list()
        if idx > 0:
            dscores_idx.extend([
                [item.item(), _idx] for item, _idx in zip(
                    cosine_similarities[0:idx], range(0,idx)
                )
            ])
        dscores_idx.extend([
            [item.item(), _idx] for item, _idx in zip(
                cosine_similarities[idx+1:], range(idx+1,len(cosine_similarities))
            )
        ])
        yield doc, dscores_idx

def get_doc_scores_from_nm_list(nm_list: List[NewsDocx]):
    documents = [' '.join(preprocess(nm.doc)) for nm in  nm_list]
    for doc__ds_idx, nm in zip(get_tfidf(documents), nm_list):
        doc, ds_idx = doc__ds_idx
        yield nm, doc, ds_idx

import os 
if __name__ == "__main__":
    nm_list = list()
    root = r'D:\Shared\test\bmw\nb_input\docs'
    count = 0
    for f in os.listdir(root):
        if not f.endswith('.docx'):
            continue
        fp = os.path.join(root, f)
        nm_list.extend(yield_news_from_docx(fp, "{}-{}-{}".format(fp[:4], fp[4:6], fp[6:8])))
        count += 1
        if count > 10:
            break


    # nm_list = list(
    #     yield_news_from_docx('20200701 BMW.docx', '2020-07-01')
    # )
    # # for nm in nm_list:
    # #     # print(asdict(nm))
    # #     pass
    # print(nm_list[0])
    # print(nm_list[-1])

    # for nm, doc, ds_idx in get_doc_scores_from_nm_list(nm_list):
    #     print(nm.header)
    #     ds_idx.sort(reverse=True, key=lambda x: x[0])
    #     for ds, idx in ds_idx[:5]:
    #         print("\t{:<3} : {:<4} : {} : {}".format(
    #             idx, round(ds, 2), nm_list[idx].label, nm_list[idx].doc[:50]
    #         ))
